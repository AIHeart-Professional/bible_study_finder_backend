"""Groups database - Data access layer."""
from typing import Optional, List
from datetime import datetime
from src.utils.logger import get_logger
from src.utils.config_loader import load_config
from pymongo import MongoClient
from bson import ObjectId
from gridfs import GridFS
from fastapi import UploadFile
from fastapi.responses import StreamingResponse
import io
import os

class GroupsDatabase:
    """Database layer for group operations."""
    
    def __init__(self):
        """Initialize the database connection."""
        self.logger = get_logger(__name__)
        self.config = load_config()
        config_url = self.config.get('database', {}).get('url', 'mongodb://localhost:27017')
        self.db_url = os.getenv("DATABASE_URL") or config_url
        if self.db_url and self.db_url.startswith("${"):
            self.db_url = 'mongodb://localhost:27017'
        self.db_name = self.config.get('database', {}).get('name', 'bible_study_finder')
        self.client = MongoClient(self.db_url)
        self.db = self.client[self.db_name]
        self.groups_collection = self.db.groups
        self.memberships_collection = self.db.groupmemberships
        self.chats_collection = self.db.groupchats
        self.worksheets_collection = self.db.bibleworksheets
        self.roles_collection = self.db.roles
        self.permissions_collection = self.db.permissions
        self.group_roles_collection = self.db.groupRoles
        self.group_requests_collection = self.db.groupRequests
        self.fs = GridFS(self.db)
        self._initialize_database()
    
    def _initialize_database(self):
        """Initialize the database and create indexes if they don't exist."""
        try:
            self.groups_collection.create_index("leaderUserId")
            self.memberships_collection.create_index([("groupId", 1), ("userId", 1)], unique=True)
            self.memberships_collection.create_index("groupId")
            self.memberships_collection.create_index("userId")
            self.memberships_collection.create_index("groupRoleId")
            self.chats_collection.create_index("groupId")
            self.worksheets_collection.create_index("groupId")
            self.roles_collection.create_index("name", unique=True)
            self.permissions_collection.create_index("action", unique=True)
            self.group_roles_collection.create_index([("groupId", 1), ("roleName", 1)], unique=True)
            self.group_roles_collection.create_index("groupId")
            self.group_requests_collection.create_index("groupId")
            self.group_requests_collection.create_index("userId")
            self.group_requests_collection.create_index([("groupId", 1), ("userId", 1)])
            self.logger.info("Groups database initialized successfully")
        except Exception as e:
            self.logger.error(f"Error initializing groups database: {e}")
    
    async def create_group(
        self,
        name: str,
        description: str,
        leaderUserId: str,
        location: dict,
        image: Optional[str] = None
    ) -> Optional[str]:
        """Create a new group in the database."""
        try:
            group_doc = {
                'name': name,
                'description': description,
                'leaderUserId': ObjectId(leaderUserId),
                'location': location,
                'studyPlans': [],
                'meals': [],
                'createdAt': datetime.utcnow(),
                'updatedAt': datetime.utcnow()
            }
            if image:
                group_doc['image'] = image
            result = self.groups_collection.insert_one(group_doc)
            self.logger.info(f"Group created successfully with ID: {result.inserted_id}")
            return str(result.inserted_id)
        except Exception as e:
            self.logger.error(f"Error creating group: {e}")
            return None
    
    async def initialize_group(self, groupId: str) -> bool:
        """Initialize a group with empty arrays."""
        try:
            result = self.groups_collection.update_one(
                {'_id': ObjectId(groupId)},
                {
                    '$set': {
                        'studyPlans': [],
                        'meals': [],
                        'updatedAt': datetime.utcnow()
                    }
                }
            )
            return result.modified_count > 0
        except Exception as e:
            self.logger.error(f"Error initializing group: {e}")
            return False
    
    async def get_all_groups(self) -> List[dict]:
        """Get all groups from the database."""
        try:
            groups = self.groups_collection.find()
            group_list = []
            for group in groups:
                group['id'] = str(group['_id'])
                del group['_id']
                if 'leaderUserId' in group:
                    group['leaderUserId'] = str(group['leaderUserId'])
                group_list.append(group)
            return group_list
        except Exception as e:
            self.logger.error(f"Error fetching all groups: {e}")
            return []
    
    async def get_group_by_id(self, groupId: str) -> Optional[dict]:
        """Get group by ID."""
        try:
            group = self.groups_collection.find_one({'_id': ObjectId(groupId)})
            if group:
                group['id'] = str(group['_id'])
                del group['_id']
                if 'leaderUserId' in group:
                    group['leaderUserId'] = str(group['leaderUserId'])
            return group
        except Exception as e:
            self.logger.error(f"Error fetching group by ID: {e}")
            return None
    
    async def get_group_members(self, groupId: str) -> List[dict]:
        """Get all members of a group with role information."""
        try:
            memberships = self.memberships_collection.find({'groupId': ObjectId(groupId)})
            members = []
            for membership in memberships:
                membership['id'] = str(membership['_id'])
                del membership['_id']
                membership['groupId'] = str(membership['groupId'])
                membership['userId'] = str(membership['userId'])
                
                if 'groupRoleId' in membership:
                    group_role_id = str(membership['groupRoleId'])
                    membership['groupRoleId'] = group_role_id
                    group_role = await self.get_group_role_config_by_id(group_role_id)
                    if group_role:
                        membership['role'] = group_role['roleName']
                        membership['permissions'] = group_role['permissions']
                    else:
                        membership['role'] = 'member'
                        membership['permissions'] = []
                else:
                    membership['role'] = 'member'
                    membership['permissions'] = []
                
                members.append(membership)
            return members
        except Exception as e:
            self.logger.error(f"Error fetching group members: {e}", exc_info=True)
            return []
    
    async def get_group_chats(self, groupId: str) -> List[dict]:
        """Get all chat messages for a group."""
        try:
            chats = self.chats_collection.find({'groupId': ObjectId(groupId)}).sort('sentAt', 1)
            chat_list = []
            for chat in chats:
                chat['id'] = str(chat['_id'])
                del chat['_id']
                chat['groupId'] = str(chat['groupId'])
                chat['userId'] = str(chat['userId'])
                chat_list.append(chat)
            return chat_list
        except Exception as e:
            self.logger.error(f"Error fetching group chats: {e}")
            return []
    
    async def get_group_meals(self, groupId: str) -> List[dict]:
        """Get all meals for a group."""
        try:
            group = await self.get_group_by_id(groupId)
            if group and 'meals' in group:
                meals = []
                for meal in group['meals']:
                    if '_id' in meal:
                        meal['id'] = str(meal['_id'])
                        del meal['_id']
                    meals.append(meal)
                return meals
            return []
        except Exception as e:
            self.logger.error(f"Error fetching group meals: {e}")
            return []
    
    async def get_group_study_plans(self, groupId: str) -> List[dict]:
        """Get all study plans for a group."""
        try:
            group = await self.get_group_by_id(groupId)
            if group and 'studyPlans' in group:
                plans = []
                for plan in group['studyPlans']:
                    if '_id' in plan:
                        plan['id'] = str(plan['_id'])
                        del plan['_id']
                    plans.append(plan)
                return plans
            return []
        except Exception as e:
            self.logger.error(f"Error fetching group study plans: {e}")
            return []
    
    async def get_user_by_id(self, userId: str) -> Optional[dict]:
        """Get user by ID from users collection."""
        try:
            users_collection = self.db.users
            user = users_collection.find_one({'_id': ObjectId(userId)})
            if user:
                user['id'] = str(user['_id'])
                del user['_id']
            return user
        except Exception as e:
            self.logger.error(f"Error fetching user by ID: {e}")
            return None
    
    async def create_group_chat(
        self,
        groupId: str,
        userId: str,
        message: str
    ) -> Optional[str]:
        """Create a new group chat message."""
        try:
            chat_doc = {
                'groupId': ObjectId(groupId),
                'userId': ObjectId(userId),
                'message': message,
                'sentAt': datetime.utcnow()
            }
            result = self.chats_collection.insert_one(chat_doc)
            self.logger.info(f"Group chat created successfully with ID: {result.inserted_id}")
            return str(result.inserted_id)
        except Exception as e:
            self.logger.error(f"Error creating group chat: {e}")
            return None
    
    async def create_worksheet(
        self,
        groupId: str,
        title: str,
        content: str
    ) -> Optional[str]:
        """Create a new Bible worksheet."""
        try:
            worksheet_doc = {
                'groupId': ObjectId(groupId),
                'title': title,
                'content': content,
                'createdAt': datetime.utcnow(),
                'updatedAt': datetime.utcnow()
            }
            result = self.worksheets_collection.insert_one(worksheet_doc)
            self.logger.info(f"Worksheet created successfully with ID: {result.inserted_id}")
            return str(result.inserted_id)
        except Exception as e:
            self.logger.error(f"Error creating worksheet: {e}")
            return None
    
    async def get_group_worksheets(self, groupId: str) -> List[dict]:
        """Get all worksheets for a group."""
        try:
            worksheets = self.worksheets_collection.find({'groupId': ObjectId(groupId)}).sort('createdAt', -1)
            worksheet_list = []
            for worksheet in worksheets:
                worksheet['id'] = str(worksheet['_id'])
                del worksheet['_id']
                worksheet['groupId'] = str(worksheet['groupId'])
                worksheet_list.append(worksheet)
            return worksheet_list
        except Exception as e:
            self.logger.error(f"Error fetching group worksheets: {e}")
            return []
    
    async def upload_worksheet_file(
        self,
        groupId: str,
        title: str,
        file: UploadFile,
        file_type: str
    ) -> tuple[bool, str, Optional[str], Optional[str], Optional[str], Optional[str]]:
        """Upload a worksheet file to GridFS."""
        self.logger.debug(f"upload_worksheet_file called for group {groupId}")
        
        try:
            file_content = await self._read_file_content(file)
            stored_file_id = await self._store_file_in_gridfs(
                file, 
                file_content, 
                groupId, 
                file_type
            )
            
            extracted_content = await self._extract_text_from_file(
                file_content, 
                file_type
            )
            
            worksheet_id = await self._create_worksheet_entry(
                groupId, 
                title, 
                extracted_content, 
                stored_file_id,
                file_type
            )
            
            if worksheet_id:
                msg = "Worksheet uploaded successfully"
                self.logger.info(f"Worksheet uploaded: {worksheet_id}")
                return (True, msg, worksheet_id, stored_file_id, file.filename, file_type)
            else:
                return (False, "Failed to create worksheet entry", None, None, None, None)
            
        except Exception as e:
            self.logger.error(f"Error uploading worksheet file: {e}", exc_info=True)
            return (False, f"Error uploading file: {str(e)}", None, None, None, None)
    
    async def _read_file_content(self, file: UploadFile) -> bytes:
        """Read file content."""
        self.logger.debug(f"Reading file content: {file.filename}")
        return await file.read()
    
    async def _store_file_in_gridfs(
        self,
        file: UploadFile,
        file_content: bytes,
        groupId: str,
        file_type: str
    ) -> str:
        """Store file in GridFS."""
        self.logger.debug(f"Storing file in GridFS: {file.filename}")
        
        file_id = self.fs.put(
            file_content,
            filename=file.filename,
            content_type=file.content_type,
            groupId=groupId,
            file_type=file_type,
            uploaded_at=datetime.utcnow()
        )
        
        self.logger.info(f"File stored in GridFS with ID: {file_id}")
        return str(file_id)
    
    async def _extract_text_from_file(
        self,
        file_content: bytes,
        file_type: str
    ) -> str:
        """Extract text content from file."""
        self.logger.debug(f"Extracting text from {file_type} file")
        
        try:
            if file_type == 'pdf':
                return await self._extract_text_from_pdf(file_content)
            elif file_type == 'docx':
                return await self._extract_text_from_docx(file_content)
            else:
                return ""
        except Exception as e:
            self.logger.error(f"Error extracting text: {e}", exc_info=True)
            return f"[Content extraction failed: {str(e)}]"
    
    async def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """Get PDF file information."""
        self.logger.debug("Processing PDF file")
        try:
            from PyPDF2 import PdfReader
            from io import BytesIO
            
            self.logger.debug(f"File content size: {len(file_content)} bytes")
            pdf_file = BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            page_count = len(reader.pages)
            self.logger.info(f"PDF has {page_count} pages")
            
            return f"<p><strong>PDF Document</strong></p><p>This document has {page_count} page(s). The PDF will be displayed below.</p>"
        except ImportError as e:
            self.logger.error(f"Import error - PyPDF2 not installed: {e}", exc_info=True)
            return "[PDF processing failed - library not installed]"
        except Exception as e:
            self.logger.error(f"Error processing PDF: {e}", exc_info=True)
            return f"[PDF processing failed: {str(e)}]"
    
    async def _extract_text_from_docx(self, file_content: bytes) -> str:
        """Extract formatted text from DOCX file as HTML."""
        self.logger.debug("Starting DOCX text extraction with formatting")
        try:
            from docx import Document
            from io import BytesIO
            
            self.logger.debug(f"File content size: {len(file_content)} bytes")
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)
            
            html_content = self._convert_docx_to_html(doc)
            
            self.logger.info(f"Successfully extracted {len(html_content)} characters from DOCX")
            return html_content if html_content else "[Empty document]"
        except ImportError as e:
            self.logger.error(f"Import error - python-docx not installed: {e}", exc_info=True)
            return "[DOCX text extraction failed - library not installed]"
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {e}", exc_info=True)
            return f"[DOCX text extraction failed: {str(e)}]"
    
    def _convert_docx_to_html(self, doc) -> str:
        """Convert DOCX document to HTML."""
        html_parts = []
        current_list = None
        list_items = []
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                html_parts.append("<br>")
                continue
            
            para_style = paragraph.style.name.lower() if paragraph.style else ""
            is_list_item = any(x in para_style for x in ['list', 'bullet', 'number'])
            
            if is_list_item:
                list_type = 'ol' if 'number' in para_style else 'ul'
                if current_list != list_type:
                    if current_list and list_items:
                        html_parts.append(self._close_list(current_list, list_items))
                        list_items = []
                    current_list = list_type
                
                list_item = self._format_list_item(paragraph)
                list_items.append(list_item)
            else:
                if current_list and list_items:
                    html_parts.append(self._close_list(current_list, list_items))
                    list_items = []
                    current_list = None
                
                para_html = self._format_paragraph(paragraph)
                html_parts.append(para_html)
        
        if current_list and list_items:
            html_parts.append(self._close_list(current_list, list_items))
        
        for table in doc.tables:
            table_html = self._format_table(table)
            html_parts.append(table_html)
        
        return "\n".join(html_parts)
    
    def _close_list(self, list_type: str, items: list) -> str:
        """Close a list with all its items."""
        items_html = "\n".join(items)
        return f"<{list_type}>\n{items_html}\n</{list_type}>"
    
    def _format_list_item(self, paragraph) -> str:
        """Format a list item."""
        indent_level = self._get_indent_level(paragraph)
        indent_style = f"margin-left: {indent_level * 20}px;" if indent_level > 0 else ""
        content = self._format_runs(paragraph)
        return f"<li style='{indent_style}'>{content}</li>"
    
    def _format_paragraph(self, paragraph) -> str:
        """Format a paragraph to HTML."""
        text = paragraph.text
        if not text.strip():
            return "<br>"
        
        style = paragraph.style.name.lower() if paragraph.style else ""
        content = self._format_runs(paragraph)
        
        indent_level = self._get_indent_level(paragraph)
        alignment = self._get_alignment(paragraph)
        
        style_attrs = []
        if indent_level > 0:
            style_attrs.append(f"margin-left: {indent_level * 40}px")
        if alignment:
            style_attrs.append(f"text-align: {alignment}")
        
        style_str = f" style='{'; '.join(style_attrs)}'" if style_attrs else ""
        
        if "heading 1" in style:
            return f"<h1{style_str}>{content}</h1>"
        elif "heading 2" in style:
            return f"<h2{style_str}>{content}</h2>"
        elif "heading 3" in style:
            return f"<h3{style_str}>{content}</h3>"
        elif "heading 4" in style:
            return f"<h4{style_str}>{content}</h4>"
        else:
            return f"<p{style_str}>{content}</p>"
    
    def _get_indent_level(self, paragraph) -> int:
        """Get paragraph indent level."""
        try:
            if paragraph.paragraph_format.left_indent:
                indent_pt = paragraph.paragraph_format.left_indent.pt
                return int(indent_pt / 36)
        except:
            pass
        return 0
    
    def _get_alignment(self, paragraph) -> str:
        """Get paragraph alignment."""
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            alignment = paragraph.alignment
            if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                return "center"
            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                return "right"
            elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                return "justify"
        except:
            pass
        return ""
    
    def _format_runs(self, paragraph) -> str:
        """Format runs with bold, italic, underline, and other formatting."""
        formatted_text = ""
        for run in paragraph.runs:
            text = run.text.replace('\t', '&nbsp;&nbsp;&nbsp;&nbsp;')
            text = text.replace('\n', '<br>')
            
            styles = []
            if hasattr(run.font, 'color') and run.font.color and run.font.color.rgb:
                try:
                    rgb = run.font.color.rgb
                    color = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
                    styles.append(f"color: {color}")
                except:
                    pass
            
            if hasattr(run.font, 'size') and run.font.size:
                try:
                    size_pt = run.font.size.pt
                    styles.append(f"font-size: {size_pt}pt")
                except:
                    pass
            
            style_attr = f" style='{'; '.join(styles)}'" if styles else ""
            
            if run.bold and run.italic and run.underline:
                text = f"<strong><em><u{style_attr}>{text}</u></em></strong>"
            elif run.bold and run.italic:
                text = f"<strong><em{style_attr}>{text}</em></strong>"
            elif run.bold and run.underline:
                text = f"<strong><u{style_attr}>{text}</u></strong>"
            elif run.italic and run.underline:
                text = f"<em><u{style_attr}>{text}</u></em>"
            elif run.bold:
                text = f"<strong{style_attr}>{text}</strong>"
            elif run.italic:
                text = f"<em{style_attr}>{text}</em>"
            elif run.underline:
                text = f"<u{style_attr}>{text}</u>"
            elif style_attr:
                text = f"<span{style_attr}>{text}</span>"
            
            formatted_text += text
        return formatted_text
    
    def _format_table(self, table) -> str:
        """Format a table to HTML."""
        html = "<table border='1' style='border-collapse: collapse; width: 100%;'>"
        
        for row in table.rows:
            html += "<tr>"
            for cell in row.cells:
                cell_text = cell.text.strip()
                html += f"<td style='padding: 8px; border: 1px solid #ddd;'>{cell_text}</td>"
            html += "</tr>"
        
        html += "</table>"
        return html
    
    async def _create_worksheet_entry(
        self,
        groupId: str,
        title: str,
        content: str,
        file_id: str,
        file_type: str = 'html'
    ) -> Optional[str]:
        """Create worksheet entry in database."""
        self.logger.debug("Creating worksheet entry in database")
        
        try:
            worksheet_doc = {
                'groupId': ObjectId(groupId),
                'title': title,
                'content': content,
                'fileId': file_id,
                'contentType': file_type,
                'createdAt': datetime.utcnow(),
                'updatedAt': datetime.utcnow()
            }
            
            result = self.worksheets_collection.insert_one(worksheet_doc)
            return str(result.inserted_id)
        except Exception as e:
            self.logger.error(f"Error creating worksheet entry: {e}")
            return None
    
    async def get_file_from_gridfs(self, file_id: str) -> Optional[StreamingResponse]:
        """Get file from GridFS."""
        self.logger.debug(f"Getting file from GridFS: {file_id}")
        try:
            grid_out = self.fs.get(ObjectId(file_id))
            
            filename = grid_out.filename
            content_type = grid_out.content_type
            
            file_content = grid_out.read()
            
            self.logger.info(f"Retrieved file: {filename}, size: {len(file_content)} bytes")
            
            return StreamingResponse(
                io.BytesIO(file_content),
                media_type=content_type,
                headers={
                    "Content-Disposition": f"attachment; filename={filename}"
                }
            )
        except Exception as e:
            self.logger.error(f"Error getting file from GridFS: {e}", exc_info=True)
            return None
    
    async def join_group(
        self,
        groupId: str,
        userId: str
    ) -> bool:
        """Join a group by creating a membership record."""
        try:
            existing = self.memberships_collection.find_one({
                'groupId': ObjectId(groupId),
                'userId': ObjectId(userId)
            })
            
            if existing:
                self.logger.warning(f"User {userId} is already a member of group {groupId}")
                return False
            
            group_role = await self._get_or_create_member_role(groupId)
            
            membership_doc = {
                'groupId': ObjectId(groupId),
                'userId': ObjectId(userId),
                'groupRoleId': ObjectId(group_role['id']),
                'joinedAt': datetime.utcnow()
            }
            result = self.memberships_collection.insert_one(membership_doc)
            self.logger.info(f"User {userId} joined group {groupId} with membership ID: {result.inserted_id}")
            return True
        except Exception as e:
            self.logger.error(f"Error joining group: {e}", exc_info=True)
            return False
    
    async def leave_group(
        self,
        groupId: str,
        userId: str
    ) -> bool:
        """Leave a group by removing the membership record."""
        try:
            membership_result = self.memberships_collection.delete_one({
                'groupId': ObjectId(groupId),
                'userId': ObjectId(userId)
            })
            
            if membership_result.deleted_count > 0:
                self.logger.info(f"User {userId} left group {groupId}")
                return True
            else:
                self.logger.warning(f"User {userId} is not a member of group {groupId}")
                return False
        except Exception as e:
            self.logger.error(f"Error leaving group: {e}")
            return False
    
    async def get_groups_by_user_id(self, userId: str) -> List[dict]:
        """Get all groups that a user is a member of."""
        try:
            # Find all memberships for this user
            memberships = self.memberships_collection.find({'userId': ObjectId(userId)})
            group_ids = [membership['groupId'] for membership in memberships]
            
            if not group_ids:
                return []
            
            # Get all groups for these group IDs
            groups = self.groups_collection.find({'_id': {'$in': group_ids}})
            group_list = []
            for group in groups:
                group['id'] = str(group['_id'])
                del group['_id']
                if 'leaderUserId' in group:
                    group['leaderUserId'] = str(group['leaderUserId'])
                group_list.append(group)
            
            return group_list
        except Exception as e:
            self.logger.error(f"Error fetching groups by user ID: {e}")
            return []
    
    # Role Management Methods
    async def create_permission(self, action: str, description: str) -> Optional[str]:
        """Create a new permission."""
        try:
            permission_doc = {
                'action': action,
                'description': description
            }
            result = self.permissions_collection.insert_one(permission_doc)
            self.logger.info(f"Permission created successfully with ID: {result.inserted_id}")
            return str(result.inserted_id)
        except Exception as e:
            self.logger.error(f"Error creating permission: {e}")
            return None
    
    async def get_all_permissions(self) -> List[dict]:
        """Get all permissions."""
        try:
            permissions = self.permissions_collection.find()
            permission_list = []
            for permission in permissions:
                permission['id'] = str(permission['_id'])
                del permission['_id']
                permission_list.append(permission)
            return permission_list
        except Exception as e:
            self.logger.error(f"Error fetching permissions: {e}")
            return []
    
    async def get_permission_by_id(self, permissionId: str) -> Optional[dict]:
        """Get a permission by ID."""
        try:
            permission = self.permissions_collection.find_one({'_id': ObjectId(permissionId)})
            if permission:
                permission['id'] = str(permission['_id'])
                del permission['_id']
            return permission
        except Exception as e:
            self.logger.error(f"Error fetching permission by ID: {e}")
            return None
    
    async def update_permission(
        self,
        permissionId: str,
        action: Optional[str] = None,
        description: Optional[str] = None
    ) -> bool:
        """Update a permission."""
        try:
            update_data = {}
            if action is not None:
                update_data['action'] = action
            if description is not None:
                update_data['description'] = description
            
            if not update_data:
                return False
            
            result = self.permissions_collection.update_one(
                {'_id': ObjectId(permissionId)},
                {'$set': update_data}
            )
            
            if result.modified_count > 0:
                self.logger.info(f"Updated permission {permissionId}")
                return True
            else:
                self.logger.warning(f"Permission {permissionId} not found or no changes")
                return False
        except Exception as e:
            self.logger.error(f"Error updating permission: {e}")
            return False
    
    async def create_role(self, name: str, permissions: List[str]) -> Optional[str]:
        """Create a new role."""
        try:
            role_doc = {
                'name': name,
                'permissions': permissions
            }
            result = self.roles_collection.insert_one(role_doc)
            self.logger.info(f"Role created successfully with ID: {result.inserted_id}")
            return str(result.inserted_id)
        except Exception as e:
            self.logger.error(f"Error creating role: {e}")
            return None
    
    async def get_all_roles(self) -> List[dict]:
        """Get all roles."""
        try:
            roles = self.roles_collection.find()
            role_list = []
            for role in roles:
                role['id'] = str(role['_id'])
                del role['_id']
                role_list.append(role)
            return role_list
        except Exception as e:
            self.logger.error(f"Error fetching roles: {e}")
            return []
    
    async def get_role_by_id(self, roleId: str) -> Optional[dict]:
        """Get a role by ID."""
        try:
            role = self.roles_collection.find_one({'_id': ObjectId(roleId)})
            if role:
                role['id'] = str(role['_id'])
                del role['_id']
            return role
        except Exception as e:
            self.logger.error(f"Error fetching role by ID: {e}")
            return None
    
    async def update_role(
        self,
        roleId: str,
        name: Optional[str] = None,
        permissions: Optional[List[str]] = None
    ) -> bool:
        """Update a role."""
        try:
            update_data = {}
            if name is not None:
                update_data['name'] = name
            if permissions is not None:
                update_data['permissions'] = permissions
            
            if not update_data:
                return False
            
            result = self.roles_collection.update_one(
                {'_id': ObjectId(roleId)},
                {'$set': update_data}
            )
            
            if result.modified_count > 0:
                self.logger.info(f"Updated role {roleId}")
                return True
            else:
                self.logger.warning(f"Role {roleId} not found or no changes")
                return False
        except Exception as e:
            self.logger.error(f"Error updating role: {e}")
            return False
    
    async def get_role_by_name(self, name: str) -> Optional[dict]:
        """Get a role by name."""
        try:
            role = self.roles_collection.find_one({'name': name})
            if role:
                role['id'] = str(role['_id'])
                del role['_id']
            return role
        except Exception as e:
            self.logger.error(f"Error fetching role by name: {e}")
            return None
    
    async def create_group_role(
        self,
        userId: str,
        groupId: str,
        role: str
    ) -> Optional[str]:
        """Assign a role to a user in a group by updating their membership."""
        try:
            self.logger.debug(f"create_group_role called with userId={userId}, groupId={groupId}, role={role}")
            
            group_role_config = await self.get_group_role_config_by_name(groupId, role)
            if not group_role_config:
                self.logger.warning(f"Group role config not found for role '{role}' in group {groupId}, creating default")
                base_role = await self.get_role_by_name(role)
                if not base_role:
                    self.logger.error(f"Base role '{role}' does not exist")
                    return None
                
                group_role_id = await self.create_group_role_config(
                    groupId=groupId,
                    roleName=role,
                    permissions=base_role.get('permissions', [])
                )
                if not group_role_id:
                    return None
                group_role_config = await self.get_group_role_config_by_id(group_role_id)
            
            existing = self.memberships_collection.find_one({
                'groupId': ObjectId(groupId),
                'userId': ObjectId(userId)
            })
            
            if existing:
                result = self.memberships_collection.update_one(
                    {
                        'groupId': ObjectId(groupId),
                        'userId': ObjectId(userId)
                    },
                    {
                        '$set': {'groupRoleId': ObjectId(group_role_config['id'])}
                    }
                )
                if result.modified_count > 0 or result.matched_count > 0:
                    self.logger.info(f"Updated membership role for user {userId} in group {groupId} to '{role}'")
                    return str(existing['_id'])
                return None
            else:
                membership_doc = {
                    'userId': ObjectId(userId),
                    'groupId': ObjectId(groupId),
                    'groupRoleId': ObjectId(group_role_config['id']),
                    'joinedAt': datetime.utcnow()
                }
                result = self.memberships_collection.insert_one(membership_doc)
                self.logger.info(f"Membership created with role '{role}' for user {userId} in group {groupId}, ID: {result.inserted_id}")
                return str(result.inserted_id)
        except Exception as e:
            self.logger.error(f"Error creating group role: {e}", exc_info=True)
            return None
    
    async def update_group_role(
        self,
        userId: str,
        groupId: str,
        role: str
    ) -> bool:
        """Update a group role by updating the membership."""
        try:
            self.logger.debug(f"update_group_role called with userId={userId}, groupId={groupId}, role={role}")
            
            group_role_config = await self.get_group_role_config_by_name(groupId, role)
            if not group_role_config:
                self.logger.warning(f"Group role config not found for role '{role}' in group {groupId}, creating default")
                base_role = await self.get_role_by_name(role)
                if not base_role:
                    self.logger.error(f"Base role '{role}' does not exist")
                    return False
                
                group_role_id = await self.create_group_role_config(
                    groupId=groupId,
                    roleName=role,
                    permissions=base_role.get('permissions', [])
                )
                if not group_role_id:
                    return False
                group_role_config = await self.get_group_role_config_by_id(group_role_id)
            
            result = self.memberships_collection.update_one(
                {
                    'groupId': ObjectId(groupId),
                    'userId': ObjectId(userId)
                },
                {
                    '$set': {'groupRoleId': ObjectId(group_role_config['id'])}
                }
            )
            
            if result.modified_count > 0:
                self.logger.info(f"Updated membership role for user {userId} in group {groupId} to {role}")
                return True
            else:
                self.logger.warning(f"Membership not found for user {userId} in group {groupId}")
                return False
        except Exception as e:
            self.logger.error(f"Error updating group role: {e}", exc_info=True)
            return False
    
    async def get_group_roles(self, groupId: Optional[str] = None, userId: Optional[str] = None) -> List[dict]:
        """Get group roles from memberships with role details, optionally filtered by groupId or userId."""
        try:
            self.logger.debug(f"get_group_roles called with groupId={groupId}, userId={userId}")
            
            query = {}
            if groupId:
                query['groupId'] = ObjectId(groupId)
            if userId:
                query['userId'] = ObjectId(userId)
            
            memberships = self.memberships_collection.find(query)
            role_list = []
            for membership in memberships:
                role_data = {
                    'id': str(membership['_id']),
                    'userId': str(membership['userId']),
                    'groupId': str(membership['groupId']),
                }
                
                if 'groupRoleId' in membership:
                    group_role = await self.get_group_role_config_by_id(str(membership['groupRoleId']))
                    if group_role:
                        role_data['role'] = group_role['roleName']
                        role_data['permissions'] = group_role['permissions']
                    else:
                        role_data['role'] = 'member'
                        role_data['permissions'] = []
                else:
                    role_data['role'] = 'member'
                    role_data['permissions'] = []
                
                role_list.append(role_data)
            
            self.logger.debug(f"Found {len(role_list)} group role(s)")
            return role_list
        except Exception as e:
            self.logger.error(f"Error fetching group roles: {e}", exc_info=True)
            return []
    
    async def remove_group_role(self, userId: str, groupId: str) -> bool:
        """Remove a role from a user in a group by setting role back to 'member'."""
        try:
            self.logger.debug(f"remove_group_role called with userId={userId}, groupId={groupId}")
            
            member_role = await self._get_or_create_member_role(groupId)
            
            result = self.memberships_collection.update_one(
                {
                    'groupId': ObjectId(groupId),
                    'userId': ObjectId(userId)
                },
                {
                    '$set': {'groupRoleId': ObjectId(member_role['id'])}
                }
            )
            
            if result.modified_count > 0:
                self.logger.info(f"Reset role to 'member' for user {userId} in group {groupId}")
                return True
            else:
                self.logger.warning(f"Membership not found for user {userId} in group {groupId}")
                return False
        except Exception as e:
            self.logger.error(f"Error removing group role: {e}", exc_info=True)
            return False
    
    async def get_user_role_in_group(self, groupId: str, userId: str) -> Optional[str]:
        """Get a user's role in a specific group from memberships collection."""
        try:
            self.logger.debug(f"get_user_role_in_group called with groupId={groupId}, userId={userId}")
            
            membership = self.memberships_collection.find_one({
                'groupId': ObjectId(groupId),
                'userId': ObjectId(userId)
            })
            
            if membership:
                if 'groupRoleId' in membership:
                    group_role = await self.get_group_role_config_by_id(str(membership['groupRoleId']))
                    if group_role:
                        role = group_role['roleName']
                        self.logger.debug(f"Found role '{role}' for user {userId} in group {groupId}")
                        return role
                
                role = 'member'
                self.logger.debug(f"Found role '{role}' for user {userId} in group {groupId}")
                return role
            
            group = await self.get_group_by_id(groupId)
            if group and str(group.get('leaderUserId')) == userId:
                self.logger.debug(f"User {userId} is the leader of group {groupId}")
                return 'leader'
            
            self.logger.debug(f"No role found for user {userId} in group {groupId}")
            return None
        except Exception as e:
            self.logger.error(f"Error getting user role in group: {e}", exc_info=True)
            return None
    
    async def remove_permission(self, permissionId: str) -> bool:
        """Remove a permission."""
        try:
            result = self.permissions_collection.delete_one({'_id': ObjectId(permissionId)})
            
            if result.deleted_count > 0:
                self.logger.info(f"Removed permission {permissionId}")
                return True
            else:
                self.logger.warning(f"Permission {permissionId} not found")
                return False
        except Exception as e:
            self.logger.error(f"Error removing permission: {e}")
            return False
    
    async def remove_role(self, roleId: str) -> bool:
        """Remove a role."""
        try:
            result = self.roles_collection.delete_one({'_id': ObjectId(roleId)})
            
            if result.deleted_count > 0:
                self.logger.info(f"Removed role {roleId}")
                return True
            else:
                self.logger.warning(f"Role {roleId} not found")
                return False
        except Exception as e:
            self.logger.error(f"Error removing role: {e}")
            return False
    
    async def remove_role_from_group(self, groupId: str, role: str) -> int:
        """Remove all group roles matching a specific role name from a group by resetting to 'member'."""
        try:
            self.logger.debug(f"remove_role_from_group called with groupId={groupId}, role={role}")
            
            role_to_remove = await self.get_group_role_config_by_name(groupId, role)
            if not role_to_remove:
                self.logger.warning(f"Role '{role}' not found in group {groupId}")
                return 0
            
            member_role = await self._get_or_create_member_role(groupId)
            
            result = self.memberships_collection.update_many(
                {
                    'groupId': ObjectId(groupId),
                    'groupRoleId': ObjectId(role_to_remove['id'])
                },
                {
                    '$set': {'groupRoleId': ObjectId(member_role['id'])}
                }
            )
            
            if result.modified_count > 0:
                self.logger.info(f"Reset {result.modified_count} membership role(s) from '{role}' to 'member' in group {groupId}")
            else:
                self.logger.warning(f"No memberships found with role '{role}' in group {groupId}")
            
            return result.modified_count
        except Exception as e:
            self.logger.error(f"Error removing role from group: {e}", exc_info=True)
            return 0
    
    async def create_group_request(
        self,
        groupId: str,
        userId: str,
        requestMessage: str
    ) -> Optional[str]:
        """Create a new group request."""
        try:
            self.logger.debug(f"create_group_request called with groupId={groupId}, userId={userId}, requestMessage={requestMessage}")
            
            # Check if request already exists
            existing = self.group_requests_collection.find_one({
                'groupId': ObjectId(groupId),
                'userId': ObjectId(userId),
                'status': 'pending'
            })
            
            if existing:
                self.logger.warning(f"Pending request already exists for user {userId} in group {groupId}")
                return None
            
            request_doc = {
                'groupId': ObjectId(groupId),
                'userId': ObjectId(userId),
                'requestMessage': requestMessage,
                'status': 'pending',
                'createdAt': datetime.utcnow()
            }
            result = self.group_requests_collection.insert_one(request_doc)
            request_id = str(result.inserted_id)
            self.logger.info(f"Group request created successfully with ID: {request_id}")
            return request_id
        except Exception as e:
            self.logger.error(f"Error creating group request: {e}", exc_info=True)
            return None
    
    async def get_group_requests(self, groupId: str) -> List[dict]:
        """Get all requests for a group."""
        try:
            self.logger.debug(f"get_group_requests called with groupId={groupId}")
            
            requests = self.group_requests_collection.find({'groupId': ObjectId(groupId)}).sort('createdAt', -1)
            request_list = []
            users_collection = self.db.users
            
            for request in requests:
                request['id'] = str(request['_id'])
                del request['_id']
                request['groupId'] = str(request['groupId'])
                request['userId'] = str(request['userId'])
                
                # Fetch username for the user
                try:
                    user = users_collection.find_one({'_id': ObjectId(request['userId'])})
                    if user:
                        request['username'] = user.get('username', 'Unknown')
                    else:
                        request['username'] = 'Unknown'
                        self.logger.warning(f"User {request['userId']} not found for group request")
                except Exception as e:
                    self.logger.warning(f"Error fetching username for user {request['userId']}: {e}")
                    request['username'] = 'Unknown'
                
                request_list.append(request)
            
            self.logger.debug(f"Found {len(request_list)} group request(s) for group {groupId}")
            return request_list
        except Exception as e:
            self.logger.error(f"Error fetching group requests: {e}", exc_info=True)
            return []
    
    async def create_group_role_config(
        self,
        groupId: str,
        roleName: str,
        permissions: List[str]
    ) -> Optional[str]:
        """Create a group-specific role configuration."""
        try:
            self.logger.debug(f"create_group_role_config called with groupId={groupId}, roleName={roleName}")
            
            group_role_doc = {
                'groupId': ObjectId(groupId),
                'roleName': roleName,
                'permissions': permissions,
                'createdAt': datetime.utcnow(),
                'updatedAt': datetime.utcnow()
            }
            result = self.group_roles_collection.insert_one(group_role_doc)
            group_role_id = str(result.inserted_id)
            self.logger.info(f"Group role config created successfully with ID: {group_role_id}")
            return group_role_id
        except Exception as e:
            self.logger.error(f"Error creating group role config: {e}", exc_info=True)
            return None
    
    async def get_group_role_configs(self, groupId: str) -> List[dict]:
        """Get all role configurations for a group."""
        try:
            self.logger.debug(f"get_group_role_configs called with groupId={groupId}")
            
            group_roles = self.group_roles_collection.find({'groupId': ObjectId(groupId)})
            role_list = []
            for role in group_roles:
                role['id'] = str(role['_id'])
                del role['_id']
                role['groupId'] = str(role['groupId'])
                role_list.append(role)
            
            self.logger.debug(f"Found {len(role_list)} group role config(s)")
            return role_list
        except Exception as e:
            self.logger.error(f"Error fetching group role configs: {e}", exc_info=True)
            return []
    
    async def get_group_role_config_by_name(
        self,
        groupId: str,
        roleName: str
    ) -> Optional[dict]:
        """Get a group role config by group ID and role name."""
        try:
            self.logger.debug(f"get_group_role_config_by_name called with groupId={groupId}, roleName={roleName}")
            
            role = self.group_roles_collection.find_one({
                'groupId': ObjectId(groupId),
                'roleName': roleName
            })
            if role:
                role['id'] = str(role['_id'])
                del role['_id']
                role['groupId'] = str(role['groupId'])
            return role
        except Exception as e:
            self.logger.error(f"Error fetching group role config by name: {e}", exc_info=True)
            return None
    
    async def get_group_role_config_by_id(self, groupRoleId: str) -> Optional[dict]:
        """Get a group role config by ID."""
        try:
            self.logger.debug(f"get_group_role_config_by_id called with groupRoleId={groupRoleId}")
            
            role = self.group_roles_collection.find_one({'_id': ObjectId(groupRoleId)})
            if role:
                role['id'] = str(role['_id'])
                del role['_id']
                role['groupId'] = str(role['groupId'])
            return role
        except Exception as e:
            self.logger.error(f"Error fetching group role config by ID: {e}", exc_info=True)
            return None
    
    async def update_group_role_config(
        self,
        groupId: str,
        roleName: str,
        permissions: List[str]
    ) -> bool:
        """Update a group role configuration."""
        try:
            self.logger.debug(f"update_group_role_config called with groupId={groupId}, roleName={roleName}")
            
            result = self.group_roles_collection.update_one(
                {
                    'groupId': ObjectId(groupId),
                    'roleName': roleName
                },
                {
                    '$set': {
                        'permissions': permissions,
                        'updatedAt': datetime.utcnow()
                    }
                }
            )
            
            if result.modified_count > 0:
                self.logger.info(f"Updated group role config for {roleName} in group {groupId}")
                return True
            else:
                self.logger.warning(f"Group role config not found or no changes")
                return False
        except Exception as e:
            self.logger.error(f"Error updating group role config: {e}", exc_info=True)
            return False
    
    async def delete_group_role_config(self, groupId: str, roleName: str) -> bool:
        """Delete a group role configuration."""
        try:
            self.logger.debug(f"delete_group_role_config called with groupId={groupId}, roleName={roleName}")
            
            result = self.group_roles_collection.delete_one({
                'groupId': ObjectId(groupId),
                'roleName': roleName
            })
            
            if result.deleted_count > 0:
                self.logger.info(f"Deleted group role config for {roleName} in group {groupId}")
                return True
            else:
                self.logger.warning(f"Group role config not found")
                return False
        except Exception as e:
            self.logger.error(f"Error deleting group role config: {e}", exc_info=True)
            return False
    
    async def _get_or_create_member_role(self, groupId: str) -> dict:
        """Get or create default 'member' role for a group."""
        try:
            existing_role = await self.get_group_role_config_by_name(groupId, 'member')
            if existing_role:
                return existing_role
            
            role_id = await self.create_group_role_config(
                groupId=groupId,
                roleName='member',
                permissions=[]
            )
            
            if role_id:
                return {'id': role_id, 'roleName': 'member', 'permissions': []}
            else:
                raise Exception("Failed to create default member role")
        except Exception as e:
            self.logger.error(f"Error getting or creating member role: {e}", exc_info=True)
            raise
    
    async def get_membership_with_role(self, groupId: str, userId: str) -> Optional[dict]:
        """Get membership info with role details."""
        try:
            self.logger.debug(f"get_membership_with_role called with groupId={groupId}, userId={userId}")
            
            membership = self.memberships_collection.find_one({
                'groupId': ObjectId(groupId),
                'userId': ObjectId(userId)
            })
            
            if not membership:
                return None
            
            membership['id'] = str(membership['_id'])
            del membership['_id']
            membership['groupId'] = str(membership['groupId'])
            membership['userId'] = str(membership['userId'])
            
            if 'groupRoleId' in membership:
                group_role_id = str(membership['groupRoleId'])
                membership['groupRoleId'] = group_role_id
                group_role = await self.get_group_role_config_by_id(group_role_id)
                if group_role:
                    membership['role'] = group_role['roleName']
                    membership['permissions'] = group_role['permissions']
                else:
                    membership['role'] = 'member'
                    membership['permissions'] = []
            else:
                membership['role'] = 'member'
                membership['permissions'] = []
            
            return membership
        except Exception as e:
            self.logger.error(f"Error getting membership with role: {e}", exc_info=True)
            return None
    
    async def create_worksheet_text(
        self,
        groupId: str,
        title: str,
        content: str
    ) -> str:
        """Create a worksheet with HTML/text content (no file)."""
        self.logger.debug(f"create_worksheet_text called with groupId={groupId}, title={title}")
        
        try:
            worksheet_doc = {
                'groupId': ObjectId(groupId),
                'title': title,
                'content': content,
                'fileId': None,  # No file for text worksheets
                'contentType': 'html',  # Mark as HTML content
                'createdAt': datetime.utcnow(),
                'updatedAt': datetime.utcnow()
            }
            
            self.logger.debug("Inserting worksheet into database")
            result = self.worksheets_collection.insert_one(worksheet_doc)
            worksheet_id = str(result.inserted_id)
            
            self.logger.info(f"Worksheet created successfully with ID: {worksheet_id}")
            return worksheet_id
            
        except Exception as e:
            self.logger.error(f"Error creating worksheet: {e}", exc_info=True)
            return ""

